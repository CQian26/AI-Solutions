#!/usr/bin/env python3
"""
attachment_scanner.py
=====================
Given a local file path, extract its text content so the clause extractor can
scan it. Handles PDF, DOCX, and plain-text; unknown types return "".

Dependencies (install any you need):
    pip install pypdf python-docx pdfminer.six

- pdfminer.six is preferred for PDFs when installed (best text fidelity).
- pypdf is the fallback.
- python-docx handles .docx (not .doc).
"""

from __future__ import annotations

import logging
import re
from pathlib import Path
from typing import Optional

log = logging.getLogger(__name__)


def _read_pdf(path: Path) -> str:
    # Try pdfminer.six first (better layout handling for gov docs).
    try:
        from pdfminer.high_level import extract_text  # type: ignore
        return extract_text(str(path)) or ""
    except ImportError:
        pass
    except Exception as e:
        log.warning("pdfminer failed on %s: %s (falling back to pypdf)", path.name, e)

    try:
        from pypdf import PdfReader  # type: ignore
        reader = PdfReader(str(path))
        parts = []
        for page in reader.pages:
            try:
                parts.append(page.extract_text() or "")
            except Exception as e:  # per-page failure shouldn't kill the doc
                log.warning("pypdf page failed on %s: %s", path.name, e)
        return "\n".join(parts)
    except ImportError:
        raise RuntimeError(
            "Neither pdfminer.six nor pypdf is installed. "
            "Run: pip install pdfminer.six pypdf"
        )


def _read_docx(path: Path) -> str:
    try:
        from docx import Document  # type: ignore
    except ImportError:
        raise RuntimeError("python-docx is not installed. Run: pip install python-docx")
    doc = Document(str(path))
    parts: list[str] = []
    for p in doc.paragraphs:
        if p.text:
            parts.append(p.text)
    # Also pick up table cells (many contract forms are tables).
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                t = cell.text.strip()
                if t:
                    parts.append(t)
    return "\n".join(parts)


def _read_text(path: Path) -> str:
    for enc in ("utf-8", "utf-16", "cp1252", "latin-1"):
        try:
            return path.read_text(encoding=enc)
        except UnicodeDecodeError:
            continue
    return path.read_bytes().decode("utf-8", errors="replace")


def extract_text_from_file(path: Path) -> str:
    """Return the best-effort plain-text content of `path`.

    Returns "" (not raise) for unsupported types so callers can move on.
    """
    suf = path.suffix.lower()
    if suf == ".pdf":
        return _read_pdf(path)
    if suf == ".docx":
        return _read_docx(path)
    if suf in (".txt", ".htm", ".html", ".md", ".rtf"):
        return _strip_html(_read_text(path)) if suf in (".htm", ".html") else _read_text(path)
    return ""


def _strip_html(html: str) -> str:
    return re.sub(r"<[^>]+>", " ", html)
